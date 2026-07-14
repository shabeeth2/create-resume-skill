#!/usr/bin/env python3
"""
convert.py - Convert a Markdown resume to DOCX, HTML, or PDF.

Part of the create-resume agent skill.
https://github.com/shabeeth2/create-resume

Usage:
  python convert.py <resume.md>              # -> resume.docx (default)
  python convert.py <resume.md> --html       # -> resume.docx + resume.html
  python convert.py <resume.md> --pdf        # -> resume.docx + resume.html + resume.pdf
  python convert.py <resume.md> --html-only  # -> resume.html only (original behavior fallback)
  python convert.py --version                # -> print version and exit
  python convert.py --help                   # -> print this help

Requirements:
  markdown-it-py  (pip install markdown-it-py)
  python-docx     (pip install python-docx)
"""

__version__ = "1.2.0"

import sys
import re
import subprocess
from pathlib import Path

# ─── HTML/PDF Styling CSS ─────────────────────────────────────────────────────
RESUME_CSS = """
@page {
  size: A4;
  margin: 14.5mm 11.9mm;
}

* {
  box-sizing: border-box;
}

body {
  margin: 0;
  padding: 0;
  background: white;
  color: black;
  font-family: 'Times New Roman', Georgia, serif;
  font-size: 12px;
  line-height: 1.3;
  text-align: left;
  -webkit-hyphens: auto;
  hyphens: auto;
}

#resume {
  max-width: 794px;
  margin: 0 auto;
}

/* Headings */
h1, h2, h3 {
  font-weight: bold;
}

h1 {
  font-size: 2.5em;
  letter-spacing: 0.1em;
  text-align: center;
  margin-bottom: 0.25em;
  border-bottom: 1px solid darkgrey;
}

h2 {
  font-size: 1.2em;
  margin-bottom: 0.25em;
  margin-top: 1em;
  border-bottom: 1px solid darkgrey;
}

h3 {
  font-size: 1.1em;
  margin-bottom: 0.25em;
  margin-top: 0.8em;
}

p, li, dl {
  margin: 0;
  margin-bottom: 5px;
}

ul, ol {
  padding-left: 1.5em;
  margin: 0.2em 0 1em 0;
}

ul { list-style-type: disc; }
ol { list-style-type: decimal; }

dl {
  display: flex;
  margin: 0;
  margin-bottom: 2px;
}

dl dt,
dl dd:not(:last-child) {
  flex: 1;
}

dl dd {
  margin: 0;
}

a {
  color: black;
  text-decoration: none;
}

img {
  max-width: 100%;
}

@media print {
  body { background: white !important; color: black !important; }
  a    { color: black !important; }
}
"""

# ─── Markdown pre-processor ───────────────────────────────────────────────────

def parse_frontmatter(text: str) -> tuple[str, dict]:
    """Parse and remove YAML frontmatter, returning (body, metadata)."""
    text = text.strip()
    metadata = {}
    if text.startswith("---"):
        end = text.find("---", 3)
        if end != -1:
            fm_text = text[3:end].strip()
            body = text[end + 3:].lstrip("\n")
            for line in fm_text.split("\n"):
                line = line.strip()
                if ":" in line:
                    key, val = line.split(":", 1)
                    metadata[key.strip()] = val.strip().strip('"').strip("'")
            return body, metadata
    return text, metadata


def strip_frontmatter(text: str) -> str:
    """Remove YAML frontmatter (--- ... ---)."""
    body, _ = parse_frontmatter(text)
    return body


def md_inline(text: str) -> str:
    """Convert inline markdown to HTML (bold, italic, links, code)."""
    text = re.sub(r"\*\*\*(.+?)\*\*\*", r"<strong><em>\1</em></strong>", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"\*(.+?)\*", r"<em>\1</em>", text)
    text = re.sub(r"_(.+?)_", r"<em>\1</em>", text)
    text = re.sub(r"`(.+?)`", r"<code>\1</code>", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2">\1</a>', text)
    return text


def preprocess_deflists(text: str) -> str:
    """Convert resume deflist syntax to HTML <dl> elements."""
    lines = text.split("\n")
    output = []
    i = 0

    while i < len(lines):
        line = lines[i]
        j = i + 1
        while j < len(lines) and lines[j].strip() == "":
            j += 1

        if j < len(lines) and re.match(r"^  : |^: ", lines[j]):
            term = line.strip()
            defs = []
            i = j
            while i < len(lines) and re.match(r"^  : |^: ", lines[i]):
                d = re.sub(r"^  : |^: ", "", lines[i]).strip()
                defs.append(d)
                i += 1
            dd_parts = "".join(f"<dd>{md_inline(d)}</dd>" for d in defs)
            output.append(f"<dl><dt>{md_inline(term)}</dt>{dd_parts}</dl>")
        else:
            output.append(line)
            i += 1

    return "\n".join(output)


# ─── HTML conversion ──────────────────────────────────────────────────────────

def md_to_html_body(md_text: str) -> str:
    from markdown_it import MarkdownIt
    md_text = strip_frontmatter(md_text)
    md_text = preprocess_deflists(md_text)
    md = MarkdownIt("commonmark", {"html": True})
    html = md.render(md_text)
    return html


def build_full_html(body: str, title: str = "Resume") -> str:
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{title}</title>
  <script src="https://code.iconify.design/2/2.2.1/iconify.min.js"></script>
  <style>{RESUME_CSS}</style>
</head>
<body>
<div id="resume">
{body}
</div>
</body>
</html>"""


# ─── PDF Engine ───────────────────────────────────────────────────────────────

def _find_browser() -> tuple[str | None, str]:
    import platform
    os_name = platform.system()
    candidates = []

    if os_name == "Windows":
        candidates = [
            (r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe", "Edge"),
            (r"C:\Program Files\Microsoft\Edge\Application\msedge.exe", "Edge"),
            (r"C:\Program Files\Google\Chrome\Application\chrome.exe", "Chrome"),
        ]
    elif os_name == "Darwin":
        candidates = [
            ("/Applications/Google Chrome.app/Contents/MacOS/Google Chrome", "Chrome"),
            ("/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge", "Edge"),
        ]
    else:
        for name in ("google-chrome", "google-chrome-stable", "chromium-browser", "chromium"):
            try:
                result = subprocess.run(["which", name], capture_output=True, text=True, timeout=5)
                if result.returncode == 0 and result.stdout.strip():
                    candidates.append((result.stdout.strip(), "Chrome"))
            except Exception:
                pass

    for path, label in candidates:
        if Path(path).exists():
            return path, label
    return None, ""


def html_to_pdf(html_path: Path, pdf_path: Path) -> bool:
    browser, label = _find_browser()
    if not browser:
        print("  [WARN] No supported browser found for PDF export.")
        return False

    file_url = html_path.as_uri()
    cmd = [
        browser,
        "--headless=new",
        "--disable-gpu",
        "--no-sandbox",
        "--run-all-compositor-stages-before-draw",
        "--virtual-time-budget=5000",
        f"--print-to-pdf={pdf_path}",
        file_url,
    ]
    print(f"  [*] Using {label} headless -> {pdf_path.name} ...")
    try:
        subprocess.run(cmd, capture_output=True, timeout=30)
    except Exception as e:
        print(f"  [WARN] Browser error: {e}")
        return False

    if pdf_path.exists() and pdf_path.stat().st_size > 0:
        return True

    cwd_pdf = Path.cwd() / pdf_path.name
    if cwd_pdf.exists() and cwd_pdf != pdf_path:
        cwd_pdf.rename(pdf_path)
        return True
    return False


# ─── DOCX Conversion ──────────────────────────────────────────────────────────

def clean_md_inline_tags(text: str) -> str:
    """Remove Markdown/HTML formatting tags for raw text runs."""
    text = re.sub(r"<span[^>]*>", "", text)
    text = re.sub(r"</span>", "", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\*\*\*(.+?)\*\*\*", r"\1", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def md_to_docx(md_text: str, docx_path: Path, font_name: str = 'Times New Roman') -> None:
    """Parse Markdown resume to produce a beautifully styled Word Document."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn

    doc = Document()

    # Set page margins (A4 matching the template CSS)
    # Set page margins (narrower for single-page fit)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)
        section.page_width = Inches(8.27) # A4
        section.page_height = Inches(11.69)

    # Base style setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = font_name
    font.size = Pt(8.5)
    font.color.rgb = RGBColor(0, 0, 0)

    md_text = strip_frontmatter(md_text)
    lines = md_text.split('\n')

    i = 0
    in_header = True
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # Header 1 (Full Name)
        if line.startswith('# '):
            name = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(name)
            run.font.name = font_name
            run.font.size = Pt(18)
            run.bold = True
            
            # Subtitle border bottom line
            pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                             r'<w:bottom w:val="single" w:sz="6" w:space="4" w:color="A9A9A9"/>'
                             r'</w:pBdr>')
            p._p.get_or_add_pPr().append(pBdr)
            i += 1
            continue

        # Header 2 (Sections)
        if line.startswith('## '):
            in_header = False
            section_title = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(section_title)
            run.font.name = font_name
            run.font.size = Pt(10.5)
            run.bold = True
            
            # Bottom border for section
            pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                             r'<w:bottom w:val="single" w:sz="6" w:space="4" w:color="A9A9A9"/>'
                             r'</w:pBdr>')
            p._p.get_or_add_pPr().append(pBdr)
            i += 1
            continue

        # Header 3 (Subsections)
        if line.startswith('### '):
            title = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(title)
            run.font.name = font_name
            run.font.size = Pt(9.0)
            run.bold = True
            i += 1
            continue

        # Lookahead: Check for three-column definition list structure
        # e.g.,
        # term
        #   : def1
        #   : def2
        next_idx = i + 1
        while next_idx < len(lines) and not lines[next_idx].strip():
            next_idx += 1
        
        if next_idx < len(lines) and re.match(r"^  : |^: ", lines[next_idx]):
            # Collect elements
            term = line
            defs = []
            i = next_idx
            while i < len(lines) and (re.match(r"^  : |^: ", lines[i]) or not lines[i].strip()):
                if lines[i].strip():
                    val = re.sub(r"^  : |^: ", "", lines[i]).strip()
                    defs.append(val)
                i += 1
            
            col_contents = [clean_md_inline_tags(term).strip()] + [clean_md_inline_tags(d).strip() for d in defs]
            
            if in_header:
                # Format contact lines centered with pipe separator (matching Image 2)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(1.5)
                p.paragraph_format.space_before = Pt(0)
                
                # Combine elements with '  |  '
                run = p.add_run("  |  ".join(col_contents))
                run.font.name = font_name
                run.font.size = Pt(8.5)
            else:
                # Build three-column table for alignment (removing borders)
                num_cols = len(col_contents)
                table = doc.add_table(rows=1, cols=num_cols)
                table.autofit = False
                
                # Clean borders
                tblPr = table._tbl.tblPr
                tblBorders = parse_xml(r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                       r'<w:top w:val="none"/><w:left w:val="none"/>'
                                       r'<w:bottom w:val="none"/><w:right w:val="none"/>'
                                       r'<w:insideH w:val="none"/><w:insideV w:val="none"/>'
                                       r'</w:tblBorders>')
                tblPr.append(tblBorders)

                row = table.rows[0]
                # Distribute widths (total page width = 7.47 inches with 0.4 margins)
                total_width = Inches(7.47)
                col_width = total_width / num_cols
                
                for idx, content in enumerate(col_contents):
                    cell = row.cells[idx]
                    cell.width = col_width
                    cell.paragraphs[0].text = "" # Clear default para
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(1.5)
                    p.paragraph_format.space_before = Pt(1.5)
                    
                    # Right align the last column (e.g. date)
                    if idx == num_cols - 1:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif idx == 1 and num_cols == 3:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    run = p.add_run(content)
                    run.font.name = font_name
                    run.font.size = Pt(8.5)
                    # Bold job titles / terms
                    if idx == 0:
                        run.bold = True
            continue

        # Bullet List Items
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.left_indent = Inches(0.2)
            
            # Basic inline parsing for bolding
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(clean_md_inline_tags(part))
                run.font.name = font_name
                run.font.size = Pt(8.5)
            i += 1
            continue

        # Numbered List Items
        if re.match(r'^\d+\.\s', line):
            match = re.match(r'^(\d+\.)\s(.*)', line)
            num_prefix = match.group(1)
            content = match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1.5)
            
            run_num = p.add_run(num_prefix + " ")
            run_num.bold = True
            run_num.font.name = font_name
            run_num.font.size = Pt(8.5)
            
            run_text = p.add_run(clean_md_inline_tags(content))
            run_text.font.name = font_name
            run_text.font.size = Pt(8.5)
            i += 1
            continue

        # Standard Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        
        # Check if line contains inline styling
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(clean_md_inline_tags(part))
            run.font.name = font_name
            run.font.size = Pt(8.5)
        
        # Center align contact lines if they slip through as plain paragraphs
        if '<span class="iconify"' in line or 'href' in line or '@' in line:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        i += 1

    doc.save(docx_path)


# ─── Command CLI Execution ───────────────────────────────────────────────────

def convert(md_file: str, docx_only: bool = True, html: bool = False, pdf: bool = False) -> None:
    md_path = Path(md_file).resolve()
    if not md_path.exists():
        print(f"[ERROR] File not found: {md_path}")
        sys.exit(1)

    docx_path = md_path.with_suffix(".docx")
    html_path = md_path.with_suffix(".html")
    pdf_path  = md_path.with_suffix(".pdf")

    print(f"\n[*] Input  : {md_path}")

    # Generate DOCX (default output)
    if docx_only or html or pdf:
        print("[1] Converting Markdown to DOCX ...")
        md_text = md_path.read_text(encoding="utf-8")
        _, metadata = parse_frontmatter(md_text)
        font_name = metadata.get('font', 'Times New Roman')
        md_to_docx(md_text, docx_path, font_name=font_name)
        print(f"    [OK] DOCX saved -> {docx_path}")

    # Generate HTML if requested
    if html or pdf:
        print("[2] Converting Markdown to HTML ...")
        md_text = md_path.read_text(encoding="utf-8")
        body_html = md_to_html_body(md_text)
        full_html = build_full_html(body_html, title=md_path.stem.replace("-", " ").title())
        html_path.write_text(full_html, encoding="utf-8")
        print(f"    [OK] HTML saved -> {html_path}")

    # Generate PDF if requested
    if pdf:
        print("[3] Converting HTML to PDF ...")
        ok = html_to_pdf(html_path, pdf_path)
        if ok:
            size_kb = pdf_path.stat().st_size // 1024
            print(f"    [OK] PDF saved  -> {pdf_path}  ({size_kb} KB)")
        else:
            print(f"    [WARN] PDF not generated.")
            print(f"    [INFO] Open '{html_path.name}' in browser -> Ctrl+P -> Save as PDF")

    print()


if __name__ == "__main__":
    args = sys.argv[1:]

    if not args or "--help" in args or "-h" in args:
        print(__doc__)
        sys.exit(0)

    if "--version" in args or "-v" in args:
        print(f"convert.py {__version__} (create-resume skill)")
        sys.exit(0)

    # Parsed arguments
    pdf = "--pdf" in args
    html = "--html" in args or pdf
    html_only = "--html-only" in args
    
    md_files = [a for a in args if not a.startswith("--")]

    if not md_files:
        print("[ERROR] Please provide a .md file path.")
        sys.exit(1)

    if html_only:
        # Fallback to the original HTML-only generation behavior
        convert(md_files[0], docx_only=False, html=True, pdf=False)
    else:
        convert(md_files[0], docx_only=True, html=html, pdf=pdf)
